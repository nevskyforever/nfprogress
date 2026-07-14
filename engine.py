import math
import os
import pickle
import platform
import sys
import tempfile
from datetime import datetime, timedelta, date, date as date_type
from pathlib import Path
from collections import defaultdict
from docx import Document

# Режим разработчика
dev_mode = "__compiled__" not in globals()

# Версия приложения
version = '4.5.2'

# Определяем систему

SYSTEM = platform.system()  # 'Windows', 'Darwin' (macOS), 'Linux'

def get_app_data_dir():
    """
    Возвращает путь к директории для хранения данных приложения
    в зависимости от операционной системы
    """
    if SYSTEM == 'Windows':
        # Windows: C:\Users\<USER>\Documents\MyAppData
        base_dir = Path(os.environ.get('USERPROFILE', '')) / 'Documents'
    elif SYSTEM == 'Darwin':  # macOS
        # macOS: /Users/<USER>/Documents/MyAppData
        base_dir = Path.home() / 'Documents'
    else:  # Linux и другие
        # Linux: /home/<USER>/Documents или ~/.local/share/MyApp
        base_dir = Path.home() / 'Documents'
        # Если папки Documents нет, используем стандартную директорию для данных
        if not base_dir.exists():
            base_dir = Path.home() / '.local' / 'share'

    # Создаем папку для нашего приложения
    app_data_dir = base_dir / 'nfprogress'

    # Создаем директорию, если её нет
    try:
        app_data_dir.mkdir(parents=True, exist_ok=True)
    except PermissionError:
        # Если нет прав на запись в Documents, используем домашнюю папку
        app_data_dir = Path.home() / '.nfprogress'
        app_data_dir.mkdir(parents=True, exist_ok=True)

    return app_data_dir


def get_test_data_dir():
    """Возвращает путь к директории тестовых данных и создаёт её при необходимости."""
    test_dir = get_app_data_dir() / 'test_data'
    test_dir.mkdir(parents=True, exist_ok=True)
    return test_dir


def sync_test_data():
    """Копирует рабочие файлы данных в папку test_data, если они новее.

    Для каждого *.pkl файла из основной директории данных создаёт
    (или перезаписывает) одноимённый файл в test_data, но только если
    рабочий файл отсутствует в test_data или новее уже лежащей там копии
    (по времени последнего изменения). Более свежие файлы в test_data
    не затираются устаревшими рабочими файлами. Файлы, которых нет
    в основной директории, не создаются и не удаляются в test_data.
    """
    import shutil
    source_dir = get_app_data_dir()
    test_dir = get_test_data_dir()
    for data_file in source_dir.glob('*.pkl'):
        test_file = test_dir / data_file.name
        if not test_file.exists() or data_file.stat().st_mtime > test_file.stat().st_mtime:
            shutil.copy2(data_file, test_file)


def get_data_file_path(name):
    """Возвращает путь к файлу данных.

    В режиме разработчика файлы читаются и пишутся из папки test_data.
    """
    if dev_mode:
        return get_test_data_dir() / f'{name}.pkl'
    return get_app_data_dir() / f'{name}.pkl'


def resource_path(relative_path):
    """Получить путь к ресурсу, работает и в .py, и в .app, и в .exe"""
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller создает временную папку и хранит путь в _MEIPASS
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)


def now_for_test():
    """Возвращает текущие дату и время с учетом режима разработчика."""
    # Тестовая дата должна работать только в настоящем режиме разработчика.
    # Если этот флаг случайно попал в пользовательские данные, обычная сборка
    # не должна навсегда застревать на выбранном вручную дне.
    settings = load_settings()
    if dev_mode and settings.get('today_for_test_mode', False):
        selected_datetime = settings.get('today_for_test_datetime')
        if isinstance(selected_datetime, datetime):
            return selected_datetime

    return datetime.combine(date.today(), datetime.now().time())


def today_for_test():
    """Возвращает сегодняшнюю дату."""
    return now_for_test().date()


def is_developer_test_date_enabled():
    """Возвращает True, если дата вручную задана в режиме разработчика."""
    if not dev_mode:
        return False
    return load_settings().get('today_for_test_mode', False)


STREAK_FREEZE_MARKER = 'freeze'
STREAK_DATE_BUG_START = date(2026, 7, 9)
STREAK_DATE_BUG_LAST_SAFE_DAY = STREAK_DATE_BUG_START - timedelta(days=1)


def is_streak_date(value):
    return isinstance(value, date) and not isinstance(value, datetime)


def iter_streak_days(streaks):
    """Возвращает пары (запись, эффективная дата) для дат и маркеров заморозки."""
    current_day = None
    if not isinstance(streaks, list):
        return

    for entry in streaks:
        if is_streak_date(entry):
            current_day = entry
            yield entry, current_day
        elif entry == STREAK_FREEZE_MARKER and current_day is not None:
            current_day = current_day + timedelta(days=1)
            yield entry, current_day


def streak_last_day(streaks):
    last_day = None
    for _, effective_day in iter_streak_days(streaks):
        last_day = effective_day
    return last_day


def streak_contains_day(streaks, target_day):
    return any(effective_day == target_day for _, effective_day in iter_streak_days(streaks))


def streak_length(streaks):
    return sum(1 for _ in iter_streak_days(streaks))


def streak_is_freeze_for_day(streaks, target_day):
    return any(entry == STREAK_FREEZE_MARKER and effective_day == target_day
               for entry, effective_day in iter_streak_days(streaks))


def remove_streak_day(streaks, target_day):
    if not isinstance(streaks, list):
        return False

    current_day = None
    for index, entry in enumerate(streaks):
        if is_streak_date(entry):
            current_day = entry
        elif entry == STREAK_FREEZE_MARKER and current_day is not None:
            current_day = current_day + timedelta(days=1)
        else:
            continue

        if current_day == target_day:
            del streaks[index]
            return True
    return False


def repair_streak_gap_after_date_bug(streaks, target_day):
    """Заполняет разрыв, возникший из-за зависшей смены даты после 09.07.2026."""
    if is_developer_test_date_enabled():
        return False
    if not isinstance(streaks, list):
        return False

    last_day = streak_last_day(streaks)
    if not last_day or last_day < STREAK_DATE_BUG_LAST_SAFE_DAY:
        return False
    if last_day >= target_day:
        return False

    changed = False
    current_day = last_day + timedelta(days=1)
    while current_day <= target_day:
        if not streak_contains_day(streaks, current_day):
            streaks.append(current_day)
            changed = True
        current_day += timedelta(days=1)
    return changed


def rebuild_streak_ending_at(end_day, length):
    """Создает непрерывный стрик заданной длины с последним днем end_day."""
    if not isinstance(end_day, date) or length <= 0:
        return []

    start_day = end_day - timedelta(days=length - 1)
    return [start_day + timedelta(days=offset) for offset in range(length)]

def load_settings():
    """Загружает данные из кроссплатформенной директории"""
    data_file = get_data_file_path('settings')
    try:
        with open(data_file, 'rb') as f:
            return pickle.load(f)
    except (FileNotFoundError, EOFError):
        # Если файл не найден, создаём пустую структуру
        return {
            'game_mode': False,
            'inf_project': False,
            'global_streak': False,
            'show_written_today_in_all_projects': False,
            'notification_display_time': 10,
            'inventory_filter': 'Все',
        }

def save_settings(data):
    """Сохраняет данные в кроссплатформенную директорию"""
    data_file = get_data_file_path('settings')
    atomic_pickle_save(data, data_file)


def atomic_pickle_save(data, data_file):
    """Атомарно сохраняет pickle в файл через уникенный временный файл в той же директории."""
    data_file = Path(data_file)
    data_file.parent.mkdir(parents=True, exist_ok=True)

    fd, temp_path_str = tempfile.mkstemp(
        prefix=f'{data_file.stem}_',
        suffix='.tmp',
        dir=str(data_file.parent)
    )
    temp_path = Path(temp_path_str)

    try:
        with os.fdopen(fd, 'wb') as f:
            pickle.dump(data, f)
            f.flush()
            os.fsync(f.fileno())
        temp_path.replace(data_file)
    except Exception:
        if temp_path.exists():
            temp_path.unlink()
        raise

class Project:
    def __init__(self, name='Без имени', goal=None,
                 create_date=None, total_symbols=0, progress=0,
                 notes=None, streaks=None, max_streak=None, streak_status='No', deadline='Нет',
                 status='активен', unit='symbols', personal_goal_for_the_day=0):

        self._name = name
        self._goal = goal  # хранится в выбранной единице
        self.create_date = create_date if create_date else today_for_test()
        self.edit_date = create_date if create_date else today_for_test()
        self.complete_date = None
        self._total_symbols = total_symbols  # хранится в выбранной единице
        self._progress = progress
        self._deadline = deadline
        self._status = status
        self.notes = notes if notes else []
        self.streaks = streaks if streaks else []
        self.max_streak = max_streak if max_streak else 0
        self.streak_status = streak_status
        self.unit = unit
        self.synch = None
        self.last_synch = None
        self.last_streak_bonus = None
        self.last_streak_lost_date = None
        self.freezes = 0
        self.deadline_set_date = today_for_test()
        self.personal_goal_for_the_day = personal_goal_for_the_day
        self.project_plan = {}

    def migrate(self):
        """Проверяет наличие всех атрибутов и добавляет недостающие"""
        defaults = {
            '_name': 'Без имени',
            '_goal': None,
            'create_date': today_for_test(),
            'edit_date': today_for_test(),
            'complete_date': None,
            '_total_symbols': 0,
            '_progress': 0,
            '_deadline': 'Нет',
            '_status': 'активен',
            'notes': [],
            'streaks': [],
            'max_streak': 0,
            'streak_status': 'No',
            'unit': 'symbols',
            'synch': None,
            'last_synch': None,
            'last_streak_bonus': None,
            'last_streak_lost_date': None,
            'freezes': 0,
            'deadline_set_date': today_for_test(),
            'project_plan': {},
        }

        for attr, default_value in defaults.items():
            if not hasattr(self, attr):
                setattr(self, attr, default_value)
            elif attr in ('notes', 'streaks') and not isinstance(getattr(self, attr), list):
                setattr(self, attr, [])

        # --- Вычисляем динамические данные ПОСЛЕ того, как базовые атрибуты созданы ---
        if not hasattr(self, 'personal_goal_for_the_day'):
            self.personal_goal_for_the_day = self.get_today_goal_in_unit()

        # Если если нет плана - генерируем его
        if self.project_plan == {}:
            self.get_today_goal_value()

        if self.synch is not None and isinstance(self.synch, str):
            self.synch = {'type': 'word', 'path': self.synch}

        # deadline_set_date должен храниться как date, иначе старт плана "прыгает".
        if self.deadline_set_date is None or not isinstance(self.deadline_set_date, date):
            self.deadline_set_date = today_for_test()

    @property
    def name(self):
        return self._name

    @name.setter
    def name(self, name):
        if hasattr(self, '_name') and self._name == name:
            self._name = name
            return

        data = load_data()
        projects = data['projects']
        names = [i for i in projects.keys() if i != self]
        if name != '':
            if name in names:
                raise ValueError('Проект с таким именем уже существует!')
            else:
                self._name = name

    @property
    def goal(self):
        return self._goal

    @goal.setter
    def goal(self, goal):
        try:
            self._goal = float(goal)
        except ValueError:
            raise ValueError('Цель должна быть числом!')

    @property
    def deadline(self):
        if self._deadline != 'Нет':
            return self._deadline
        return 'Нет'

    @deadline.setter
    def deadline(self, deadline):
        if deadline == '':
            self._deadline = 'Нет'
            self.deadline_set_date = None
        else:
            self._deadline = deadline
            self.deadline_set_date = today_for_test()

    @property
    def deadline_str(self):
        if self._deadline != 'Нет':
            return datetime.strftime(self.deadline, '%d.%m.%y')
        return 'Нет'

    @property
    def status(self):
        return self._status

    @status.setter
    def status(self, status):
        self._status = status

    @property
    def total_units(self):
        return self._total_symbols

    @total_units.setter
    def total_units(self, total_symbols):
        self._total_symbols = total_symbols

    def get_goal_symbols(self):
        """Возвращает цель в символах."""
        if self._goal == float('inf'):
            return float('inf')
        return unit_converter(self.unit, self._goal, 'symbols')

    def get_total_symbols(self):
        """Возвращает текущее количество в символах."""
        return unit_converter(self.unit, self._total_symbols, 'symbols')

    @property
    def progress(self):
        """Прогресс проекта в процентах"""
        goal_sym = self.get_goal_symbols()
        total_sym = self.get_total_symbols()
        if goal_sym and goal_sym > 0 and goal_sym != float('inf'):
            self._progress = total_sym / goal_sym * 100
        else:
            self._progress = 0
        return self._progress

    def get_added_symbols_today_value(self):
        """Возвращает количество добавленных сегодня символов."""
        today = today_for_test()
        today_added = [i.get_added_symbols() for i in self.notes if i.get_date_create() == today]
        return sum(today_added) if today_added else 0

    def get_added_today_in_unit(self):
        """Возвращает количество добавленных сегодня в единице проекта."""
        added_sym = self.get_added_symbols_today_value()
        return unit_converter('symbols', added_sym, self.unit)

    def get_today_goal_value(self):
        """
        Возвращает накопительный план на сегодня в символах.
        Генерирует план от текущей даты до дедлайна по принципу "лесенки".
        """
        today = today_for_test()

        if getattr(self, 'project_plan', None) is None:
            self.project_plan = {}

        plan = self.project_plan

        # === 1. ПРОВЕРКА ДЕДЛАЙНА (Если он наступил или прошел) ===
        if self.deadline != 'Нет' and isinstance(self.deadline, date):
            if today >= self.deadline:
                return self.get_goal_symbols()

        # === 2. БЫСТРЫЙ ВОЗВРАТ ИЗ КЭША (Защита от "убегания" цели) ===
        # Создаем "слепок" параметров, от которых зависит план
        current_signature = str((
            getattr(self, 'deadline', 'Нет'),
            getattr(self, 'personal_goal_for_the_day', 0),
            getattr(self, '_goal', 0),
            getattr(self, 'unit', 'symbols'),
            getattr(self, 'deadline_set_date', None)
        ))

        # Если день есть в плане И настройки не менялись с момента последнего расчета — отдаем кэш
        if today in plan and plan.get('signature') == current_signature:
            return plan[today]

        # === 3. ПЕРЕСЧЕТ ПЛАНА (Инициализация) ===

        today_added_symbols = 0
        if hasattr(self, 'streaks') and isinstance(self.streaks, list):
            for streak in self.streaks:
                if isinstance(streak, dict) and streak.get('date') == today:
                    today_added_symbols = streak.get('count', 0)
                    break

        base_total = unit_converter(self.unit, self.total_units) - today_added_symbols
        daily_goal_symbols = 0

        if getattr(self, 'personal_goal_for_the_day', 0):
            converted = unit_converter(self.unit, self.personal_goal_for_the_day, 'symbols')
            if converted == float('inf'):
                daily_goal_symbols = float('inf')
            else:
                # Убираем ceil – оставляем дробное значение
                daily_goal_symbols = converted

        elif self.deadline != 'Нет' and isinstance(self.deadline, date):

            if self.goal != float('inf'):
                goal_sym = unit_converter(self.unit, self.goal)
                total_days = (self.deadline - today).days + 1
                if total_days > 0:
                    remaining = max(0, goal_sym - base_total)
                    # Деление без округления вверх
                    daily_goal_symbols = remaining / total_days
                else:
                    daily_goal_symbols = goal_sym

        # === 4. ГЕНЕРАЦИЯ ЛЕСЕНКИ ПЛАНА ===

        # Если кэш устарел (настройки поменялись) или его нет — строим заново
        if not plan or plan.get('signature') != current_signature:
            plan.clear()
            plan['signature'] = current_signature  # Сохраняем слепок настроек

            end_date = self.deadline if (self.deadline != 'Нет' and isinstance(self.deadline, date)) else (
                        today + timedelta(days=30))

            current_date = today
            current_goal = base_total

            while current_date <= end_date:
                current_goal += daily_goal_symbols
                if self.deadline != 'Нет' and isinstance(self.deadline, date):
                    max_sym = self.get_goal_symbols()
                    if current_goal > max_sym:
                        current_goal = max_sym

                plan[current_date] = current_goal
                current_date += timedelta(days=1)

        else:
            # Сценарий достройки плана (если закончились 30 дней запаса)
            # Отфильтровываем строковые ключи (наш signature), оставляем только даты
            date_keys = [k for k in plan.keys() if isinstance(k, date)]
            if not date_keys:
                last_date = today
                current_goal = base_total
            else:
                last_date = max(date_keys)
                current_goal = plan[last_date]

            current_date = last_date + timedelta(days=1)
            end_date = today + timedelta(days=30)

            while current_date <= end_date:
                current_goal += daily_goal_symbols
                if self.deadline != 'Нет' and isinstance(self.deadline, date):
                    max_sym = self.get_goal_symbols()
                    if current_goal > max_sym:
                        current_goal = max_sym
                plan[current_date] = current_goal
                current_date += timedelta(days=1)

        return plan.get(today, base_total + daily_goal_symbols)

    def get_today_goal_in_unit(self):
        """Возвращает цель на сегодня в единице проекта."""
        goal_sym = self.get_today_goal_value()
        if goal_sym == float('inf'):
            return float('inf')
        return unit_converter('symbols', goal_sym, self.unit)

    def get_need_write_value(self):
        """Возвращает остаток написать в символах."""
        goal_sym = self.get_goal_symbols()
        total_sym = self.get_total_symbols()
        if goal_sym == float('inf'):
            return float('inf')
        return goal_sym - total_sym

    def get_need_write_in_unit(self):
        """Возвращает остаток написать в единице проекта."""
        need_sym = self.get_need_write_value()
        if need_sym == float('inf'):
            return float('inf')
        return unit_converter('symbols', need_sym, self.unit)

    def get_streak_status(self):
        """Возвращает статус локального стрика проекта (статусы потери только в день потери)."""
        today = today_for_test()
        yesterday = today - timedelta(days=1)
        total = self.get_total_symbols()
        planned = self.get_today_goal_value()

        # Проверяем, есть ли дедлайн
        if self.deadline == 'Нет':
            return 'No'

        # Если стрик сегодня уже продлен - не проверяем
        # Добавлена проверка на наличие элементов в self.streaks перед обращением к [-1]
        last_streak_day = streak_last_day(self.streaks)
        if self.streak_status == 'Go' and last_streak_day == today:
            return 'Go'

        # Если стрик уже завершен - дальше не проверяем
        if self.status == 'завершен':
            return 'Complete'

        # Заморозка
        if self.streak_status == 'Freeze' and streak_is_freeze_for_day(self.streaks, today):
            return 'Freeze'

        # Если проект бесконечный - стрика нет
        if planned == 0 or self.goal == float('inf'):
            return 'No'

        # Убедимся, что streaks — список
        if not isinstance(self.streaks, list):
            self.streaks = []

        day_completed = total >= planned

        if (not self.streaks and
                isinstance(self.streak_status, str) and
                self.streak_status.startswith('Lose ') and
                self.last_streak_lost_date is not None and
                self.last_streak_lost_date >= STREAK_DATE_BUG_START):
            status_parts = self.streak_status.split()
            if len(status_parts) >= 2 and status_parts[1].isdigit():
                lost_len = int(status_parts[1])
                self.streaks = rebuild_streak_ending_at(STREAK_DATE_BUG_LAST_SAFE_DAY, lost_len)
                self.streak_status = 'Active'
                self.last_streak_lost_date = None

        # === СБРОС УСТАРЕВШЕГО СТАТУСА ПОТЕРИ (чистого) ===
        if (isinstance(self.streak_status, str) and
                self.streak_status.startswith('Lose ') and
                len(self.streak_status.split()) == 2 and
                self.last_streak_lost_date is not None and
                self.last_streak_lost_date < today):
            # Прошёл день потери без активности — сбрасываем
            self.streak_status = 'No'
            self.last_streak_lost_date = None

        # === НОРМАЛИЗАЦИЯ КОМБИНИРОВАННОГО СТАТУСА (только для сегодня) ===
        if (isinstance(self.streak_status, str) and
                self.streak_status.startswith('Lose ') and
                len(self.streak_status.split()) == 3):
            # Если это не сегодняшний стрик, превращаем в обычный Start
            if streak_last_day(self.streaks) != today:
                self.streak_status = 'Start'
            else:
                parts = self.streak_status.split()
                self.streak_status = parts[2]  # сохраняем базовый статус ('Start')
                return f'Lose {parts[1]} {parts[2]}'  # возвращаем комбо для сообщения

        # === ПЕРСИСТЕНЦИЯ СТАТУСА ПОТЕРИ (только в день потери) ===
        if (isinstance(self.streak_status, str) and
                self.streak_status.startswith('Lose ') and
                len(self.streak_status.split()) == 2 and
                self.last_streak_lost_date == today and
                not day_completed):
            return self.streak_status

        # === ❗️ РЕТРОСПЕКТИВНАЯ ПРОВЕРКА ПРОПУЩЕННЫХ ДНЕЙ ===
        # Проверяем и "дописываем" прошлые дни ДО обработки сегодняшнего результата
        last_streak_day = streak_last_day(self.streaks)
        repair_streak_gap_after_date_bug(self.streaks, yesterday)
        last_streak_day = streak_last_day(self.streaks)
        if last_streak_day and last_streak_day < yesterday:
            current_day = last_streak_day + timedelta(days=1)
            plan = self.project_plan

            # Проверяем дни вплоть до вчерашнего
            while current_day < today:
                # Пропущенный день нельзя закрывать текущим накопительным total:
                # иначе написанное позже ретроактивно продлевает стрик без заморозки.
                if not apply_project_freeze(self, current_day):
                    break
                current_day += timedelta(days=1)

        # Случай 1: сегодня уже есть запись в streaks (уже продлили сегодня)
        last_streak_day = streak_last_day(self.streaks)
        if last_streak_day == today and day_completed:
            self.streak_status = 'Go' if streak_length(self.streaks) > 1 else 'Start'
            self.last_streak_lost_date = None  # стрик активен, потери нет
            return self.streak_status

        # Случай 2: день выполнен
        if day_completed:
            if not self.streaks:
                # Первый день стрика
                self.streaks = [today]
                self.streak_status = 'Start'
                self.last_streak_lost_date = None
            elif last_streak_day == yesterday:
                # Продолжение стрика (в том числе если мы его восстановили шагом выше)
                self.streaks.append(today)
                self.streak_status = 'Go'
                self.last_streak_lost_date = None
            else:
                # Перерыв: потеря предыдущего + начало нового
                # (Сюда мы попадем, если ретроспективная проверка прервалась из-за нехватки символов)
                lose_len = streak_length(self.streaks)
                if lose_len > self.max_streak:
                    self.max_streak = lose_len
                self.streaks = [today]
                # Комбинированный статус (потеря + старт)
                self.streak_status = f'Lose {lose_len} Start'
                self.last_streak_lost_date = None
                return self.streak_status
        else:
            # День не выполнен
            if not self.streaks:
                self.streak_status = 'No'
            elif last_streak_day == today:
                # Уже сегодня была запись, но план не выполнен
                self.streak_status = 'Active'
            elif last_streak_day == yesterday:
                # Вчера был стрик, сегодня ещё не выполнен → активный
                self.streak_status = 'Active'
            else:
                # Потеря стрика (пропуск >1 дня и недобор символов)
                lose_len = streak_length(self.streaks)
                if lose_len > self.max_streak:
                    self.max_streak = lose_len
                self.streaks = []
                self.streak_status = f'Lose {lose_len}'
                self.last_streak_lost_date = today

        # Обновляем максимальный стрик, если текущая длина больше
        current_streak_len = streak_length(self.streaks)
        if current_streak_len > self.max_streak:
            self.max_streak = current_streak_len

        return self.streak_status

    def get_streak_status_msg(self, msg_type=None):
        """Возвращает сообщение о статусе стрика."""
        status = self.get_streak_status()
        if status == 'Start':
            if msg_type == 'min':
                return '🔥 стрик Начат'
            return f'🔥 Стрик в {self.name} начат! Отличное начало, главное - продолжать!'
        elif status == 'Freeze':
            if msg_type == 'min':
                return f'❄️ Стрик заморожен'
            return f'❄️ Стрик в {self.name} заморожен'
        elif status == 'Go':
            streaks = streak_length(self.streaks)
            if msg_type == 'min':
                return '🚀 стрик продлен'
            return f'🚀 Стрик в {self.name} продлен! Вы движетесь к цели уже {streaks} дней подряд!'
        elif status == 'Active':
            streaks = streak_length(self.streaks)
            if msg_type == 'min':
                return f'⏳ Стрик в {streaks} дн не продлен'
            return f'⏳ Стрик в {self.name} активен ({streaks} дн.), но ещё не продлён сегодня'
        elif status == 'Complete':
            streaks = streak_length(self.streaks)
            if msg_type == 'min':
                return '🎉 СТРИК ЗАВЕРШЕН'
            return f'🎉 СТРИК ЗАВЕРШЕН! Вы выполняли цель {streaks} дней подряд, потрясающе!'
        elif status == 'No':
            return f'🙃 Стрик не начат'
        elif status.split()[0] == 'Lose':
            status_parts = status.split()
            if len(status_parts) == 2:
                if msg_type == 'min':
                    return f'💔 Потеряно {status_parts[1]} д. стрика'
                return f'💔 Стрик потерян! Вы были в цели {status_parts[1]} дней подряд.'
            elif len(status_parts) == 3:
                return (f'💔 Стрик потерян! Вы были в цели {status_parts[1]} дней подряд.'
                        f'\n🔥 Вы начали новый стрик!')
        return 'Вывод статуса не работает'

    def set_new_notes(self, new_note):
        """Добавляет заметку и обновляет total_units в единице проекта."""
        self.notes.append(new_note)
        # new_note.new_total хранится в символах, конвертируем в единицу проекта
        self._total_symbols = unit_converter('symbols', new_note.new_total, self.unit)

    def get_statistic(self):
        """
        Считает и возвращает статистику по проекту в виде словаря.
        Все параметры вычисляются из существующих данных без изменения структуры.
        """

        today = today_for_test()

        # --- Вспомогательные структуры ---
        # Группируем символы по дате
        symbols_by_date = defaultdict(float)
        for note in self.notes:
            symbols_by_date[note.get_date_create()] += note.get_added_symbols()

        active_days = sorted(symbols_by_date.keys())
        total_notes = len(self.notes)
        num_active_days = len(active_days)

        # --- 1. Кол-во записей ---
        stat_notes_count = total_notes

        # --- 2. Всего написано в единице проекта ---
        stat_total_in_unit = self._total_symbols  # уже хранится в единице проекта

        # --- 3. Среднее символов в день (только по активным дням) ---
        if num_active_days > 0:
            total_sym = self.get_total_symbols()  # в символах
            stat_avg_symbols_per_active_day = round(total_sym / num_active_days)
        else:
            stat_avg_symbols_per_active_day = 0

        # --- 4. Среднее кол-во символов в записи ---
        if total_notes > 0:
            total_added = sum(note.get_added_symbols() for note in self.notes)
            stat_avg_symbols_per_note = round(total_added / total_notes)
        else:
            stat_avg_symbols_per_note = 0

        # --- 5. Среднее кол-во записей в день (по активным дням) ---
        if num_active_days > 0:
            stat_avg_notes_per_day = round(total_notes / num_active_days, 1)
        else:
            stat_avg_notes_per_day = 0

        # --- 6. Использовано заморозок ---
        stat_freezes_used = self.freezes

        # --- 7. Лучший день (максимум символов за одну дату) ---
        if symbols_by_date:
            best_date = max(symbols_by_date, key=symbols_by_date.get)
            best_value = round(symbols_by_date[best_date])
            stat_best_day = f'{best_date.strftime("%d.%m.%Y")} — {best_value} симв.'
        else:
            stat_best_day = '—'

        # --- 8. Самый продуктивный день недели ---
        DAYS_RU = {
            0: 'Понедельник', 1: 'Вторник', 2: 'Среда',
            3: 'Четверг', 4: 'Пятница', 5: 'Суббота', 6: 'Воскресенье'
        }
        if symbols_by_date:
            symbols_by_weekday = defaultdict(float)
            for d, sym in symbols_by_date.items():
                symbols_by_weekday[d.weekday()] += sym
            best_weekday = max(symbols_by_weekday, key=symbols_by_weekday.get)
            stat_best_weekday = DAYS_RU[best_weekday]
        else:
            stat_best_weekday = '—'

        # --- 9. Текущий стрик (дней) ---
        stat_current_streak = streak_length(self.streaks)

        # --- 10. Максимальный стрик ---
        # max_streak обновляется в get_streak_status, берём максимум с текущим
        stat_max_streak = max(self.max_streak, stat_current_streak)

        # --- 11. Дней с начала проекта ---
        if self.create_date:
            stat_days_since_start = (today - self.create_date).days + 1
        else:
            stat_days_since_start = 0

        # --- 12. Активных дней (было хоть одна запись) ---
        stat_active_days_count = num_active_days

        # --- 13. Процент активных дней от общего времени проекта ---
        if stat_days_since_start > 0:
            stat_active_days_percent = round(
                stat_active_days_count / stat_days_since_start * 100, 1
            )
        else:
            stat_active_days_percent = 0

        return {
            'Кол-во записей': stat_notes_count,
            'Всего написано в единице проекта': stat_total_in_unit,
            'Среднее символов в день': stat_avg_symbols_per_active_day,
            'Среднее кол-во символов в записи': stat_avg_symbols_per_note,
            'Среднее кол-во записей в день': stat_avg_notes_per_day,
            'Использовано заморозок': stat_freezes_used,
            'Лучший день': stat_best_day,
            'Самый продуктивный день недели': stat_best_weekday,
            'Текущий стрик (дней)': stat_current_streak,
            'Максимальный стрик': stat_max_streak,
            'Дней с начала проекта': stat_days_since_start,
            'Активных дней': stat_active_days_count,
            'Процент активных дней': f'{stat_active_days_percent}%',
        }
class Stage(Project):
    def __init__(self):
        super().__init__()
        self.status = 'в работе'

class Note:

    def __init__(self, new_total, added_symbols, added_progress, date_create=None):
        if date_create is None:
            now = datetime.now()
            today = today_for_test()
            self.date_create = datetime(
                year=today.year,
                month=today.month,
                day=today.day,
                hour=now.hour,
                minute=now.minute
            )
        else:
            self.date_create = date_create

        self.new_total = new_total
        self.added_symbols = added_symbols
        self.added_progress = added_progress

    def get_new_total(self):
        return self.new_total

    def get_added_symbols(self):
        return self.added_symbols

    def get_added_progress(self):
        return self.added_progress

    def get_date_create(self):
        return self.date_create.date()

    def get_date_create_str(self):
        return self.date_create.strftime('%d.%m.%Y %H:%M')

class Notification:
    def __init__(self, text, tag=None, date_create=None, status='New'):
        self.text = text
        if date_create is None:
            now = datetime.now()
            today = today_for_test()
            self.date_create = datetime(
                year=today.year,
                month=today.month,
                day=today.day,
                hour=now.hour,
                minute=now.minute
            )
        else:
            self.date_create = date_create
        self.status = status
        self.tag = tag

    def get_date_create(self):
        return self.date_create.strftime('%H:%M %d.%m.%y')

    def get_status(self):
        return self.status

    def set_status(self, status):
        self.status = status

    def get_text(self):
        return self.text

    def set_text(self, text):
        self.text = text


# === ЗАГРУЗКА И СОХРАНЕНИЕ ===

def load_data():
    """Загружает данные из кроссплатформенной директории"""
    data_file = get_data_file_path('data')
    try:
        with open(data_file, 'rb') as f:
            data = pickle.load(f)

        # Миграция всех проектов - добавление недостающих атрибутов
        projects = data.get('projects', {})
        for project_name, project in projects.items():
            if isinstance(project, Project):
                # Вызываем метод миграции для каждого проекта
                project.migrate()

        return data

    except (FileNotFoundError, EOFError):
        # Если файл не найден, создаём пустую структуру
        return {
            'last': None,
            'projects': {},
            'notifications': [],
            'global_streaks': [],
            'global_streak_status': 'No',
            'max_global_streak': 0,
            'last_global_streak_bonus': None,
            'last_global_streak_lost_date': None,
        }

def save_data(data):
    """Сохраняет данные в кроссплатформенную директорию"""
    data_file = get_data_file_path('data')
    atomic_pickle_save(data, data_file)


def apply_project_freeze(project, freeze_day=None, gamer=None, save_gamer=True):
    """Применяет заморозку к проекту и списывает её из инвентаря игрока."""
    if freeze_day is None:
        freeze_day = today_for_test()

    if not isinstance(project, Project):
        return False
    if project.deadline == 'Нет' or project.status == 'завершен':
        return False
    if project.goal == float('inf') or project.get_today_goal_value() == 0:
        return False
    if not isinstance(project.streaks, list):
        project.streaks = []

    last_streak_day = streak_last_day(project.streaks)
    if last_streak_day != freeze_day - timedelta(days=1):
        return False
    if not load_settings().get('game_mode', False):
        return False

    if gamer is None:
        import game
        gamer = game.load_game()

    items = getattr(gamer, 'items', {})
    category_items = items.get('Предметы', {})
    if not isinstance(category_items, dict):
        return False

    freeze_names = ('Заморозка', '❄️Заморозка')
    freeze_name = next((name for name in freeze_names if category_items.get(name, 0) > 0), None)
    if freeze_name is None:
        return False

    category_items[freeze_name] -= 1
    project.streaks.append(STREAK_FREEZE_MARKER)
    if freeze_day == today_for_test():
        project.streak_status = 'Freeze'
    project.freezes += 1

    current_streak_len = streak_length(project.streaks)
    if current_streak_len > project.max_streak:
        project.max_streak = current_streak_len

    if save_gamer:
        gamer.save()

    return True


def export_data_to_file(file_path):
    """Экспортирует данные в указанный файл"""
    data = load_data()
    try:
        with open(file_path, 'wb') as f:
            pickle.dump(data, f)
        return True, "Данные успешно экспортированы"
    except Exception as e:
        return False, f"Ошибка экспорта: {e}"


def import_data_from_file(file_path):
    """Импортирует данные из указанного файла"""
    try:
        with open(file_path, 'rb') as f:
            data = pickle.load(f)
        save_data(data)
        return True, "Данные успешно импортированы"
    except Exception as e:
        return False, f"Ошибка импорта: {e}"


def get_data_directory_info():
    """Возвращает информацию о директории с данными (для отладки)"""
    data_dir = get_app_data_dir()
    data_file = get_data_file_path('data')

    return {
        'system': SYSTEM,
        'data_dir': str(data_dir),
        'data_file': str(data_file),
        'data_dir_exists': data_dir.exists(),
        'data_file_exists': data_file.exists(),
        'data_file_size': data_file.stat().st_size if data_file.exists() else 0
    }

def global_streak_status(data, today=None):
    """Возвращает статус глобального стрика. Статусы потери возвращаются только в день потери."""
    if today is None:
        today = today_for_test()
    yesterday = today - timedelta(days=1)

    streaks = data.get('global_streaks', [])
    if streaks:
        # Чистим только повторные даты. Маркеры заморозки могут повторяться подряд.
        seen_dates = set()
        normalized_streaks = []
        for streak in streaks:
            if is_streak_date(streak):
                if streak in seen_dates:
                    continue
                seen_dates.add(streak)
            normalized_streaks.append(streak)
        streaks = normalized_streaks
        data['global_streaks'] = streaks
    if not isinstance(streaks, list):
        streaks = []
        data['global_streaks'] = streaks

    max_streak = data.get('max_global_streak', 0)
    prev_status = data.get('global_streak_status', 'No')
    last_lost_date = data.get('last_global_streak_lost_date')
    last_lost_len = data.get('last_global_streak_lose_len', 0)
    projects = data.get('projects', {})

    if (not streaks and
            isinstance(prev_status, str) and
            prev_status.startswith('Lose ') and
            last_lost_date is not None and
            last_lost_date >= STREAK_DATE_BUG_START):
        repair_len = last_lost_len or max_streak
        if repair_len > 0:
            streaks = rebuild_streak_ending_at(STREAK_DATE_BUG_LAST_SAFE_DAY, repair_len)
            data['global_streaks'] = streaks
            data['global_streak_status'] = 'Active'
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
            prev_status = 'Active'
            last_lost_date = None
            last_lost_len = 0

    # Заморозка
    if prev_status == 'Freeze' and streak_is_freeze_for_day(streaks, today):
        return 'Freeze'

    # Определяем, есть ли сегодня проекты, выполнившие план и есть ли активные
    has_active_today = False
    for project in projects.values():
        if isinstance(project, Project):
            # ВАЖНО: обязательно вызываем метод для актуализации статуса на текущий день
            actual_status = project.get_streak_status()

            # Проект активен или завершен
            if project.status in ['активен', 'завершен']:
                # Стрик начат/продлен
                if actual_status in ['Start', 'Go']:
                    has_active_today = True
                # Стрик завершен сегодня
                elif actual_status == 'Complete' and streak_last_day(project.streaks) == today:
                    has_active_today = True

                # Перенимаем стрик, если он активен и длиннее текущего глобального (используем copy!)
                if actual_status in ['Start', 'Go', 'Active'] and streak_length(streaks) < streak_length(project.streaks):
                    streaks = project.streaks.copy()

    repair_streak_gap_after_date_bug(streaks, yesterday)
    data['global_streaks'] = streaks

    # === ОБРАБОТКА СОХРАНЁННОГО СТАТУСА ПОТЕРИ (только в день потери) ===
    if (isinstance(prev_status, str) and
            prev_status.startswith('Lose ') and
            len(prev_status.split()) == 2 and
            last_lost_date == today):
        # Автозаморозка могла восстановить вчерашний день уже после того,
        # как потеря была сохранена. В этом случае старый Lose больше не валиден.
        if streak_last_day(streaks) == yesterday:
            status = 'Active'
            data['global_streaks'] = streaks
            data['global_streak_status'] = status
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
            data['max_global_streak'] = max(streak_length(streaks), max_streak)
            save_data(data)
            return status

        # Если сегодня нет активности — оставляем статус потери
        if not has_active_today:
            return prev_status
        else:
            # Появилась активность в день потери — переходим к комбинированному статусу
            streaks = [today]
            status = f'Lose {last_lost_len} Start'
            data['global_streaks'] = streaks
            data['global_streak_status'] = status
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
            data['max_global_streak'] = max(1, max_streak)
            save_data(data)
            return status

    # === ОБРАБОТКА КОМБИНИРОВАННОГО СТАТУСА (только в день его установки) ===
    if (isinstance(prev_status, str) and
            prev_status.startswith('Lose ') and
            len(prev_status.split()) == 3 and
            streak_last_day(streaks) == today):
        # Это комбинированный статус, установленный сегодня — возвращаем его
        return prev_status

    # Если сохранённый статус потери относится к прошлому дню, сбрасываем его в 'No'
    if (isinstance(prev_status, str) and
            prev_status.startswith('Lose ') and
            len(prev_status.split()) == 2 and
            last_lost_date is not None and
            last_lost_date < today):
        # Прошёл день потери без активности — сбрасываем в No
        data['global_streak_status'] = 'No'
        data['last_global_streak_lost_date'] = None
        data['last_global_streak_lose_len'] = 0
        save_data(data)
        prev_status = 'No'

    # === ОСНОВНАЯ ЛОГИКА РАСЧЁТА СТАТУСА ===
    if has_active_today:
        if not streaks:
            # Новый стрик
            streaks.append(today)
            if last_lost_date is not None:
                # Если была зафиксирована потеря (например, вчера), но сегодня активность — комбо
                status = f'Lose {last_lost_len} Start'
                data['last_global_streak_lost_date'] = None
                data['last_global_streak_lose_len'] = 0
            else:
                status = 'Start'
        elif streak_last_day(streaks) == yesterday:
            # Продолжение стрика
            streaks.append(today)
            status = 'Go'
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
        elif streak_last_day(streaks) == today:
            # ИСПРАВЛЕНИЕ: Стрик уже обновлён сегодня (глобально ранее, либо только что переняли от проекта).
            # Вычисляем статус вместо возврата старого prev_status.
            if streak_length(streaks) > 1:
                status = 'Go'
            else:
                if last_lost_date is not None:
                    status = f'Lose {last_lost_len} Start'
                else:
                    status = 'Start'
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
        else:
            # Пропуск дней: потеря старого стрика и начало нового
            lose_len = streak_length(streaks)
            if lose_len > max_streak:
                max_streak = lose_len
            streaks = [today]
            status = f'Lose {lose_len} Start'
            data['last_global_streak_lost_date'] = None
            data['last_global_streak_lose_len'] = 0
    else:
        # Нет активности сегодня
        if not streaks:
            status = 'No'
        elif streak_last_day(streaks) == yesterday:
            status = 'Active'
        elif streak_last_day(streaks) == today:
            # Запись за сегодня есть, но план не выполнен (находится в процессе) — статус активен
            status = 'Active'
        else:
            # Потеря стрика из-за пропуска более одного дня
            lose_len = streak_length(streaks)
            if lose_len > max_streak:
                max_streak = lose_len
            streaks.clear()
            status = f'Lose {lose_len}'
            data['last_global_streak_lost_date'] = today
            data['last_global_streak_lose_len'] = lose_len

    # Обновление максимальной длины стрика
    current_len = streak_length(streaks)
    if current_len > max_streak:
        max_streak = current_len

    data['global_streaks'] = streaks
    data['max_global_streak'] = max_streak
    data['global_streak_status'] = status
    save_data(data)
    return status


def refresh_project_streak_statuses(data):
    """Актуализирует локальные стрики проектов без UI-уведомлений."""
    projects = data.get('projects', {})
    global_streaks = data.setdefault('global_streaks', [])
    if not isinstance(global_streaks, list):
        global_streaks = []
        data['global_streaks'] = global_streaks
    changed = False
    freeze_changed = False

    for project_name, project in projects.items():
        if not isinstance(project, Project):
            continue

        old_streaks = list(project.streaks) if isinstance(project.streaks, list) else []
        old_status = project.streak_status
        old_freezes = getattr(project, 'freezes', 0)
        old_max_streak = getattr(project, 'max_streak', 0)

        project.get_streak_status()

        if getattr(project, 'freezes', 0) > old_freezes:
            freeze_changed = True
            new_freeze_days = [
                effective_day
                for entry, effective_day in iter_streak_days(project.streaks)
                if entry == STREAK_FREEZE_MARKER
            ][-(project.freezes - old_freezes):]
            for freeze_day in new_freeze_days:
                if streak_last_day(global_streaks) == freeze_day - timedelta(days=1):
                    global_streaks.append(STREAK_FREEZE_MARKER)
                    data['global_streak_status'] = 'Freeze'
                    changed = True
        if (
                project.streaks != old_streaks
                or project.streak_status != old_status
                or getattr(project, 'freezes', 0) != old_freezes
                or getattr(project, 'max_streak', 0) != old_max_streak
                or global_streaks != data.get('global_streaks', [])
        ):
            changed = True
            data['projects'][project_name] = project

    if changed:
        save_data(data)

    return {'changed': changed, 'freeze_changed': freeze_changed}


def global_streak_status_msg(data, status=None):
    """Сообщение для глобального стрика по статусу."""
    if status is None:
        status = data.get('global_streak_status', 'No')

    streaks = data.get('global_streaks', [])

    if status == 'Start':
        return '🔥 Глобальный стрик начат!'
    elif status == 'Go':
        return f'🚀 Глобальный стрик продлен! Дней подряд: {streak_length(streaks)}'
    elif status == 'Freeze':
        return f'❄️ Глобальный стрик в {streak_length(streaks)} д. заморожен'
    elif status == 'Active':
        # Активный глобальный стрик (вчера продлен, сегодня не продлен)
        return f'👀 Глобальный стрик в {streak_length(streaks)} д. не продлен сегодня.'
    elif isinstance(status, str) and status.startswith('Lose '):
        parts = status.split()
        if len(parts) == 2 and parts[1].isdigit():
            return f'💔 Глобальный стрик потерян! Было дней подряд: {parts[1]}'
        elif len(parts) == 3:
            # Комбинированное сообщение: потеря + новый старт
            return (f'💔 Глобальный стрик потерян! Было дней подряд: {parts[1]}'
                    f'\n🔥 Вы начали новый стрик!')
        return '💔 Глобальный стрик потерян!'
    elif status == 'No':
        # Если нет активных стриков, но были раньше
        last_streak_day = streak_last_day(streaks)
        if last_streak_day is not None:
            today = today_for_test()
            days_ago = (today - last_streak_day).days
            if days_ago == 1:
                return f'👀 Глобальный стрик в {streak_length(streaks)} д. не продлен.'
            else:
                return f'😴 Глобальный стрик не активен. Последний стрик был {days_ago} дней назад'
        else:
            return '😴 Глобальный стрик не начат'

    return '😴 Глобальный стрик не начат'

def unit_converter(unit, value, convert_to=None):
    """
    Конвертирует количество value из исходной единицы unit в целевую единицу convert_to.
    """
    factors = {
        'symbols': 1,
        'A4': 1800,
        'author_list': 40000,
        'ficbook_pages': 4500
    }

    if convert_to in (None, False):
        convert_to = 'symbols'

    if unit not in factors or convert_to not in factors:
        return None

    # Приводим к общему знаменателю (символам) и вычисляем результат
    symbols_value = value * factors[unit]
    result = symbols_value / factors[convert_to]

    # Логика округления в зависимости от типа единицы
    if convert_to == 'symbols':
        return result
    elif convert_to == 'author_list':
        # Для авторских листов — 1 знак после запятой
        return round(result, 1)
    else:
        # Для страниц (A4, Ficbook) оставляем округление вверх до целого
        return math.ceil(result)

def count_symbols_in_docx(filepath):
    """
    Возвращает общее количество символов (с пробелами) в тексте документа .docx.
    Учитываются абзацы и текст в таблицах.
    """

    doc = Document(filepath)
    total = 0
    for paragraph in doc.paragraphs:
        total += len(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += len(paragraph.text)
    return total

# При импорте модуля: в режиме разработчика синхронизируем test_data с текущими данными
if dev_mode:
    sync_test_data()
else:
    get_app_data_dir()
